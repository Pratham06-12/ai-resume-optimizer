"""
Resume parsing service — PDF (PyMuPDF) and DOCX (python-docx).

Design:
- Extract plain text first (reliable baseline for ATS)
- Heuristic section splitting (no ML required in Phase 1)
- Section headers matched case-insensitively (common resume headings)
"""

import re
from io import BytesIO
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document

from app.models.schemas import ParsedResume, ResumeSection

# Common resume section headers (order matters for greedy matching)
SECTION_ALIASES: dict[str, list[str]] = {
    "skills": [
        "skills",
        "technical skills",
        "core competencies",
        "technologies",
        "tech stack",
    ],
    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "work history",
    ],
    "education": [
        "education",
        "academic background",
        "qualifications",
    ],
    "projects": [
        "projects",
        "personal projects",
        "key projects",
        "selected projects",
    ],
}

# Catch-all headers that go into other_sections
OTHER_HEADERS = [
    "summary",
    "professional summary",
    "objective",
    "certifications",
    "certificates",
    "awards",
    "volunteer",
    "languages",
    "interests",
    "publications",
]


def _normalize_whitespace(text: str) -> str:
    """Collapse excessive spaces and normalize line breaks."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _split_bullets(content: str) -> list[str]:
    """
    Split section body into bullet-like lines.
    Handles •, -, *, and numbered lists.
    """
    if not content.strip():
        return []

    lines = content.split("\n")
    items: list[str] = []
    buffer: list[str] = []

    bullet_pattern = re.compile(r"^[\s]*(?:[•\-\*●○▪]|\d+[\.\)])\s+")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if buffer:
                items.append(" ".join(buffer).strip())
                buffer = []
            continue

        if bullet_pattern.match(stripped):
            if buffer:
                items.append(" ".join(buffer).strip())
                buffer = []
            cleaned = bullet_pattern.sub("", stripped).strip()
            if cleaned:
                items.append(cleaned)
        else:
            buffer.append(stripped)

    if buffer:
        items.append(" ".join(buffer).strip())

    # If no bullets detected, return non-empty lines
    if not items:
        items = [ln.strip() for ln in lines if ln.strip()]

    return items


def _find_section_boundaries(text: str) -> list[tuple[str, int, str]]:
    """
    Returns list of (canonical_key, start_index, matched_header).
    canonical_key: skills | experience | education | projects | other
    """
    lines = text.split("\n")
    boundaries: list[tuple[str, int, str]] = []
    char_offset = 0

    all_headers: list[tuple[str, str]] = []
    for key, aliases in SECTION_ALIASES.items():
        for alias in aliases:
            all_headers.append((key, alias))
    for alias in OTHER_HEADERS:
        all_headers.append(("other", alias))

    # Sort by length descending so "work experience" matches before "experience"
    all_headers.sort(key=lambda x: len(x[1]), reverse=True)

    for i, line in enumerate(lines):
        line_clean = line.strip().lower()
        # Header lines are usually short
        if len(line_clean) > 80:
            char_offset += len(line) + 1
            continue

        for key, alias in all_headers:
            # Exact line match or line ends with header (e.g. "EXPERIENCE")
            if line_clean == alias or line_clean.rstrip(":") == alias:
                boundaries.append((key, char_offset, line.strip()))
                break

        char_offset += len(line) + 1

    # Sort by document position
    boundaries.sort(key=lambda x: x[1])
    return boundaries


def _extract_sections(text: str) -> dict[str, ResumeSection]:
    """Split full text into resume sections using header detection."""
    boundaries = _find_section_boundaries(text)

    sections: dict[str, str] = {
        "skills": "",
        "experience": "",
        "education": "",
        "projects": "",
    }
    other_parts: list[tuple[str, str]] = []

    if not boundaries:
        # No headers found — put everything in experience as fallback
        sections["experience"] = text
        return {
            k: ResumeSection(title=k.title(), content=v, items=_split_bullets(v))
            for k, v in sections.items()
        }

    for idx, (key, start, header) in enumerate(boundaries):
        end = boundaries[idx + 1][1] if idx + 1 < len(boundaries) else len(text)
        body = text[start:end]
        # Remove the header line from body
        body_lines = body.split("\n", 1)
        if len(body_lines) > 1:
            body = body_lines[1]
        else:
            body = ""

        body = body.strip()
        if key == "other":
            other_parts.append((header, body))
        elif key in sections:
            if sections[key]:
                sections[key] += "\n\n" + body
            else:
                sections[key] = body

    result: dict[str, ResumeSection] = {}
    for key, content in sections.items():
        result[key] = ResumeSection(
            title=key.title(),
            content=content,
            items=_split_bullets(content),
        )

    result["_other"] = other_parts  # type: ignore
    return result


class ResumeParserService:
    """Parse resume files from bytes or path."""

    ALLOWED_EXTENSIONS = {".pdf", ".docx"}

    @classmethod
    def validate_extension(cls, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        if ext not in cls.ALLOWED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type '{ext}'. Allowed: PDF, DOCX."
            )
        return ext.lstrip(".")

    @staticmethod
    def extract_text_from_pdf(data: bytes) -> str:
        """PyMuPDF: extract text page by page."""
        parts: list[str] = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                parts.append(page.get_text("text"))
        return _normalize_whitespace("\n".join(parts))

    @staticmethod
    def extract_text_from_docx(data: bytes) -> str:
        """python-docx: paragraphs + table cells."""
        document = Document(BytesIO(data))
        parts: list[str] = []

        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)

        return _normalize_whitespace("\n".join(parts))

    def parse(
        self,
        file_bytes: bytes,
        filename: str,
    ) -> ParsedResume:
        """
        Main entry: bytes + filename → ParsedResume.
        """
        file_type = self.validate_extension(filename)
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            raw_text = self.extract_text_from_pdf(file_bytes)
        else:
            raw_text = self.extract_text_from_docx(file_bytes)

        if not raw_text or len(raw_text) < 50:
            raise ValueError(
                "Could not extract enough text from the file. "
                "Try a text-based PDF (not scanned image) or DOCX."
            )

        section_map = _extract_sections(raw_text)
        other_raw = section_map.pop("_other", [])  # type: ignore

        other_sections: list[ResumeSection] = []
        if isinstance(other_raw, list):
            for header, body in other_raw:
                other_sections.append(
                    ResumeSection(
                        title=header,
                        content=body,
                        items=_split_bullets(body),
                    )
                )

        word_count = len(raw_text.split())

        return ParsedResume(
            raw_text=raw_text,
            skills=section_map.get(
                "skills",
                ResumeSection(title="Skills", content="", items=[]),
            ),
            experience=section_map.get(
                "experience",
                ResumeSection(title="Experience", content="", items=[]),
            ),
            education=section_map.get(
                "education",
                ResumeSection(title="Education", content="", items=[]),
            ),
            projects=section_map.get(
                "projects",
                ResumeSection(title="Projects", content="", items=[]),
            ),
            other_sections=other_sections,
            word_count=word_count,
            file_name=filename,
            file_type=file_type,
        )
